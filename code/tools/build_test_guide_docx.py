"""Build the Word test guide from its canonical Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CODE_ROOT = Path(__file__).resolve().parents[1]
SOURCE = CODE_ROOT / "docs" / "课题三指标3.1-3.9算法测试大纲与运行说明.md"
TARGET = SOURCE.with_suffix(".docx")

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "5F6B76"
TABLE_FILL = "E8EEF5"
CODE_FILL = "F2F4F7"
BODY_FONT = "Noto Sans CJK SC"


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, monospace: bool = False) -> None:
    western = "Menlo" if monospace else BODY_FONT
    east_asia = BODY_FONT
    run.font.name = western
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), western)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), western)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 5,
                          line: float = 1.2) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            row_properties = row._tr.get_or_add_trPr()
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            row_properties.append(header)
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=2, line=1.15)
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, bold=row_index == 0, color=INK)


def create_numbering(document: Document, *, numbered: bool) -> int:
    numbering = document.part.numbering_part.element
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    style_name = "ListNumber" if numbered else "ListBullet"
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        paragraph_style = abstract.find(f".//{qn('w:pStyle')}")
        if paragraph_style is not None and paragraph_style.get(qn("w:val")) == style_name:
            abstract_id = int(abstract.get(qn("w:abstractNumId")))
            break
    if abstract_id is None:
        raise ValueError(f"Word 模板缺少编号样式: {style_name}")

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    if numbered:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        number.append(override)
    numbering.append(number)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    number_properties.append(level)
    number_properties.append(number)
    paragraph_properties.append(number_properties)


def add_inline(paragraph, text: str, *, default_size: float = 11) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]), size=default_size, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=default_size, bold=True, color=DARK_BLUE)
        else:
            set_run_font(paragraph.add_run(token[1:-1]), size=default_size - 0.5, color=INK, monospace=True)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]), size=default_size, color=INK)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph)
    add_inline(paragraph, text)


def add_bullet(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=4)
    apply_numbering(paragraph, num_id)
    add_inline(paragraph, text)


def add_numbered(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=4)
    apply_numbering(paragraph, num_id)
    add_inline(paragraph, text)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=3, after=7, line=1.05)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shading)
    set_run_font(paragraph.add_run("\n".join(lines)), size=8.5, color=INK, monospace=True)


def add_heading(document: Document, level: int, text: str) -> None:
    if level == 1:
        paragraph = document.add_paragraph(style="Title")
        set_paragraph_spacing(paragraph, before=0, after=12, line=1.0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(paragraph.add_run(text), size=22, bold=True, color=DARK_BLUE)
        paragraph.paragraph_format.keep_with_next = True
        return
    style_name = "Heading 1" if level == 2 else "Heading 2"
    paragraph = document.add_paragraph(style=style_name)
    paragraph.add_run(text)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    title = document.styles["Title"]
    title.font.name = BODY_FONT
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 1.0

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, after=0, line=1.0)
    set_run_font(header.add_run("课题三算法测试大纲与运行说明"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer, after=0, line=1.0)
    set_run_font(footer.add_run("第 "), size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    set_run_font(footer.add_run(" 页"), size=8.5, color=MUTED)

    bullet_num_id = create_numbering(document, numbered=False)
    decimal_num_id = create_numbering(document, numbered=True)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    previous_was_numbered = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            previous_was_numbered = False
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            previous_was_numbered = False
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = parse_table(table_lines)
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for col_index, value in enumerate(row):
                    cell = table.cell(row_index, col_index)
                    cell.text = ""
                    add_inline(cell.paragraphs[0], value, default_size=9.5)
            if len(rows[0]) == 3:
                widths = [1800, 2880, 4680]
            else:
                base = CONTENT_WIDTH_DXA // len(rows[0])
                widths = [base] * (len(rows[0]) - 1)
                widths.append(CONTENT_WIDTH_DXA - sum(widths))
            configure_table(table, widths)
            document.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            previous_was_numbered = False
            add_heading(document, len(heading.group(1)), heading.group(2))
        elif stripped.startswith("- "):
            previous_was_numbered = False
            add_bullet(document, stripped[2:], bullet_num_id)
        else:
            numbered = re.match(r"^(\d+\.)\s+(.+)$", stripped)
            if numbered:
                if not previous_was_numbered:
                    decimal_num_id = create_numbering(document, numbered=True)
                add_numbered(document, numbered.group(2), decimal_num_id)
                previous_was_numbered = True
            else:
                previous_was_numbered = False
                add_body_paragraph(document, stripped)
        index += 1

    document.core_properties.title = "课题三指标 3.1-3.9 算法测试大纲与运行说明"
    document.core_properties.subject = "算法文件、测试代码、通过判据与逐项运行说明"
    document.core_properties.author = "课题三项目组"
    document.core_properties.comments = ""
    document.core_properties.last_modified_by = "课题三项目组"
    return document


def main() -> None:
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
